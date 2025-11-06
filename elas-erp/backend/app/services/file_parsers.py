"""
File parsers for different file types (CSV, XLSX, PDF, DOCX)
"""
import os
import pandas as pd
from typing import Tuple, List, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document
import openpyxl
from io import StringIO

def parse_csv_or_txt(file_path: str) -> pd.DataFrame:
    """Parse CSV or TXT file"""
    try:
        # Try common delimiters
        for delimiter in [',', '\t', ';', '|']:
            try:
                df = pd.read_csv(file_path, delimiter=delimiter)
                if len(df.columns) > 1:  # Valid multi-column file
                    return df
            except:
                continue
        
        # Fallback: let pandas auto-detect
        return pd.read_csv(file_path)
    except Exception as e:
        raise ValueError(f"Failed to parse CSV/TXT: {str(e)}")


def parse_excel(file_path: str) -> pd.DataFrame:
    """Parse XLSX/XLS file"""
    try:
        # Read first sheet by default
        df = pd.read_excel(file_path, engine='openpyxl')
        return df
    except Exception as e:
        raise ValueError(f"Failed to parse Excel file: {str(e)}")


def parse_pdf(file_path: str) -> pd.DataFrame:
    """
    Parse PDF file and extract tables/text
    
    Strategy:
    1. Try to extract text from PDF
    2. Look for tabular data patterns
    3. Convert to DataFrame
    
    Note: This is a basic implementation. For production, consider using
    libraries like tabula-py or camelot-py for better table extraction.
    """
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Extract text from all pages
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            # Try to parse as CSV-like text
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            
            if not lines:
                raise ValueError("No text extracted from PDF")
            
            # Simple heuristic: check if lines contain common separators
            # This is basic - for real PDFs with tables, use tabula-py
            sample = lines[0] if lines else ""
            delimiter = None
            
            for sep in [',', '\t', '|', ';']:
                if sep in sample:
                    delimiter = sep
                    break
            
            if delimiter:
                # Try to parse as delimited text
                csv_text = '\n'.join(lines)
                df = pd.read_csv(StringIO(csv_text), delimiter=delimiter)
                return df
            else:
                # Create a single-column DataFrame with the text
                # User can see the extracted text in the UI
                return pd.DataFrame({
                    'extracted_text': lines[:100],  # Limit to 100 lines
                    'line_number': range(1, min(101, len(lines) + 1))
                })
                
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}. Try using a CSV/Excel export instead.")


def parse_docx(file_path: str) -> pd.DataFrame:
    """
    Parse DOCX file and extract tables
    
    Strategy:
    1. Extract all tables from document
    2. Convert first table to DataFrame
    3. If no tables, extract text paragraphs
    """
    try:
        doc = Document(file_path)
        
        # Try to extract tables first
        if doc.tables:
            # Get first table
            table = doc.tables[0]
            
            # Convert table to list of lists
            data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                data.append(row_data)
            
            if not data:
                raise ValueError("Empty table found")
            
            # First row as header
            df = pd.DataFrame(data[1:], columns=data[0])
            return df
        else:
            # No tables, extract paragraphs as text
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            if not paragraphs:
                raise ValueError("No content found in DOCX")
            
            return pd.DataFrame({
                'text': paragraphs[:100],  # Limit to 100 paragraphs
                'paragraph_number': range(1, min(101, len(paragraphs) + 1))
            })
            
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}. Try using a CSV/Excel export instead.")


def parse_file(file_path: str) -> Tuple[pd.DataFrame, str]:
    """
    Universal file parser that detects file type and routes to appropriate parser
    
    Returns:
        Tuple of (DataFrame, file_type)
    """
    file_ext = Path(file_path).suffix.lower()
    
    try:
        if file_ext in ['.csv', '.txt', '.tsv']:
            df = parse_csv_or_txt(file_path)
            return df, 'CSV/TXT'
        
        elif file_ext in ['.xlsx', '.xls']:
            df = parse_excel(file_path)
            return df, 'Excel'
        
        elif file_ext == '.pdf':
            df = parse_pdf(file_path)
            return df, 'PDF'
        
        elif file_ext == '.docx':
            df = parse_docx(file_path)
            return df, 'DOCX'
        
        else:
            raise ValueError(
                f"Unsupported file type: {file_ext}. "
                f"Supported types: CSV, TXT, TSV, XLSX, XLS, PDF, DOCX"
            )
    
    except Exception as e:
        # Re-raise with more context
        raise ValueError(f"Error parsing {file_ext} file: {str(e)}")


def validate_dataframe(df: pd.DataFrame, min_rows: int = 1, min_cols: int = 1) -> None:
    """
    Validate that DataFrame meets minimum requirements
    
    Raises ValueError if validation fails
    """
    if df.empty:
        raise ValueError("File contains no data")
    
    if len(df) < min_rows:
        raise ValueError(f"File must contain at least {min_rows} rows of data")
    
    if len(df.columns) < min_cols:
        raise ValueError(f"File must contain at least {min_cols} columns")
    
    # Check for all-null columns
    null_cols = df.columns[df.isnull().all()].tolist()
    if null_cols:
        raise ValueError(f"File contains empty columns: {', '.join(null_cols)}")

